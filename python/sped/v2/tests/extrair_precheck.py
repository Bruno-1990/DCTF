#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para extrair conteúdo do arquivo Precheck_XML_SPED_EFD.docx
"""

from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("❌ python-docx não está instalado. Instale com: pip install python-docx")
    sys.exit(1)

def main():
    # Caminho do arquivo
    docx_path = Path(r"C:\Users\bruno\Desktop\SPED 2.0\Precheck_XML_SPED_EFD.docx")
    
    if not docx_path.exists():
        print(f"❌ Arquivo não encontrado: {docx_path}")
        return
    
    try:
        print(f"📄 Lendo documento: {docx_path}")
        doc = Document(str(docx_path))
        
        # Extrair todo o texto
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extrair texto de tabelas
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        text_content = "\n".join(full_text)
        
        print(f"\n✅ Texto extraído: {len(text_content)} caracteres")
        
        # Salvar em arquivo de texto
        output_path = Path(__file__).parent / "precheck_extracted.txt"
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text_content)
        
        print(f"\n💾 Texto extraído salvo em: {output_path}")
        print(f"\n📋 Primeiros 3000 caracteres:\n")
        print("=" * 80)
        print(text_content[:3000])
        print("=" * 80)
        
    except Exception as e:
        print(f"❌ Erro ao ler documento: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
