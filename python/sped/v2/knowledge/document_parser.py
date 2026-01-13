"""
Parser de Documentos Legais (PDF/DOCX)
Extrai texto, estrutura e metadados de documentos legais para o sistema de conhecimento
"""

import os
import hashlib
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, field
from datetime import datetime, date
import json

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️  pdfplumber não disponível. Instale com: pip install pdfplumber")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx não disponível. Instale com: pip install python-docx")


@dataclass
class DocumentMetadata:
    """Metadados extraídos de um documento legal"""
    documento_tipo: str  # GUIA_PRATICO, ATO_COTEPE, CONVENIO, PORTARIA, NOTA_TECNICA
    documento_nome: str
    versao: Optional[str] = None
    vigencia_inicio: Optional[date] = None
    vigencia_fim: Optional[date] = None
    autor: Optional[str] = None
    data_publicacao: Optional[date] = None
    orgao: Optional[str] = None  # CONFAZ, SEFAZ-ES, etc.
    arquivo_path: str = ""
    hash_arquivo: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class DocumentSection:
    """Seção estruturada de um documento"""
    title: str
    level: int  # 1 = título principal, 2 = subtítulo, etc.
    content: str
    page_number: Optional[int] = None
    article_number: Optional[str] = None  # "Art. 12", "§ 3º", etc.
    subsections: List['DocumentSection'] = field(default_factory=list)


@dataclass
class ParsedDocument:
    """Documento parseado completo"""
    metadata: DocumentMetadata
    full_text: str
    sections: List[DocumentSection]
    tables: List[Dict[str, Any]] = field(default_factory=list)
    pages: List[str] = field(default_factory=list)  # Texto por página


class PDFParser:
    """Parser para documentos PDF"""
    
    def __init__(self, file_path: str):
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")
        
        if not PDF_AVAILABLE:
            raise ImportError("pdfplumber não está instalado")
    
    def calculate_hash(self) -> str:
        """Calcula hash SHA-256 do arquivo"""
        sha256_hash = hashlib.sha256()
        with open(self.file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    
    def extract_metadata(self) -> DocumentMetadata:
        """Extrai metadados do PDF e do nome do arquivo"""
        # Extrair metadados do nome do arquivo
        filename = self.file_path.stem
        metadata = PDFParser._parse_filename(filename)
        
        # Calcular hash
        metadata.hash_arquivo = self.calculate_hash()
        metadata.arquivo_path = str(self.file_path)
        
        # Tentar extrair metadados do PDF
        try:
            with pdfplumber.open(self.file_path) as pdf:
                if pdf.metadata:
                    pdf_meta = pdf.metadata
                    if not metadata.autor and pdf_meta.get('Author'):
                        metadata.autor = pdf_meta.get('Author')
                    if not metadata.data_publicacao and pdf_meta.get('CreationDate'):
                        try:
                            # Tentar parsear data do PDF
                            date_str = pdf_meta.get('CreationDate', '')
                            # Formato comum: D:20250101120000
                            if date_str.startswith('D:'):
                                date_str = date_str[2:10]  # YYYYMMDD
                                metadata.data_publicacao = datetime.strptime(date_str, '%Y%m%d').date()
                        except:
                            pass
                    
                    metadata.metadata['pdf_metadata'] = {
                        'title': pdf_meta.get('Title'),
                        'subject': pdf_meta.get('Subject'),
                        'creator': pdf_meta.get('Creator'),
                        'producer': pdf_meta.get('Producer'),
                    }
        except Exception as e:
            print(f"⚠️  Erro ao extrair metadados do PDF: {e}")
        
        return metadata
    
    @staticmethod
    def _parse_filename(filename: str) -> DocumentMetadata:
        """Parse do nome do arquivo para extrair tipo, versão, etc."""
        filename_upper = filename.upper()
        
        # Identificar tipo de documento
        documento_tipo = "OUTRO"
        documento_nome = filename
        versao = None
        vigencia_inicio = None
        vigencia_fim = None
        
        # Padrões de identificação
        if "GUIA PRATICO" in filename_upper or "GUIA_PRATICO" in filename_upper:
            documento_tipo = "GUIA_PRATICO"
            documento_nome = "Guia Prático EFD ICMS/IPI"
            # Extrair versão (ex: "3.2.1", "v3.2.1")
            versao_match = re.search(r'[vV]?(\d+\.\d+(?:\.\d+)?)', filename)
            if versao_match:
                versao = versao_match.group(1)
        
        elif "ATO COTEPE" in filename_upper or "ATO_COTEPE" in filename_upper:
            documento_tipo = "ATO_COTEPE"
            # Extrair número do ato (ex: "79_25", "79/25")
            ato_match = re.search(r'(\d+)[_/](\d+)', filename)
            if ato_match:
                versao = f"{ato_match.group(1)}/{ato_match.group(2)}"
                documento_nome = f"Ato COTEPE/ICMS {versao}"
        
        elif "CONVENIO" in filename_upper or "CONVÊNIO" in filename_upper:
            documento_tipo = "CONVENIO"
            # Extrair número do convênio (ex: "142_18", "142/18")
            conv_match = re.search(r'(\d+)[_/](\d+)', filename)
            if conv_match:
                versao = f"{conv_match.group(1)}/{conv_match.group(2)}"
                documento_nome = f"Convênio ICMS {versao}"
        
        elif "PORTARIA" in filename_upper:
            documento_tipo = "PORTARIA"
            # Extrair número da portaria
            port_match = re.search(r'(\d+)[-_]?[Rr]?', filename)
            if port_match:
                versao = port_match.group(1)
                documento_nome = f"Portaria {versao}"
        
        elif "NOTA TECNICA" in filename_upper or "NOTA_TECNICA" in filename_upper or "NT" in filename_upper:
            documento_tipo = "NOTA_TECNICA"
            # Extrair número da NT (ex: "010_2025", "06_2016")
            nt_match = re.search(r'(\d+)[_/](\d{4})', filename)
            if nt_match:
                versao = f"{nt_match.group(1)}/{nt_match.group(2)}"
                documento_nome = f"Nota Técnica {versao}"
        
        elif "MOC" in filename_upper or "MANUAL" in filename_upper:
            documento_tipo = "MANUAL"
            documento_nome = "Manual de Orientação ao Contribuinte (MOC)"
            # Extrair versão (ex: "7.0", "v7.0")
            versao_match = re.search(r'[vV]?(\d+\.\d+)', filename)
            if versao_match:
                versao = versao_match.group(1)
        
        # Tentar extrair datas de vigência do nome
        # Padrão: "01/2025" ou "2025-01"
        date_match = re.search(r'(\d{2})[/-](\d{4})', filename)
        if date_match:
            try:
                month, year = int(date_match.group(1)), int(date_match.group(2))
                vigencia_inicio = date(year, month, 1)
            except:
                pass
        
        return DocumentMetadata(
            documento_tipo=documento_tipo,
            documento_nome=documento_nome,
            versao=versao,
            vigencia_inicio=vigencia_inicio,
            vigencia_fim=vigencia_fim
        )
    
    def extract_text(self) -> ParsedDocument:
        """Extrai texto completo, estrutura e páginas do PDF"""
        metadata = self.extract_metadata()
        full_text = ""
        pages = []
        sections: List[DocumentSection] = []
        
        with pdfplumber.open(self.file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                try:
                    page_text = page.extract_text() or ""
                    pages.append(page_text)
                    full_text += f"\n\n--- Página {page_num} ---\n\n{page_text}"
                    
                    # Extrair tabelas da página
                    tables = page.extract_tables()
                    for table in tables:
                        # Processar tabela (pode ser expandido)
                        pass
                
                except Exception as e:
                    print(f"⚠️  Erro ao processar página {page_num}: {e}")
                    pages.append("")
        
        # Identificar estrutura (seções, artigos)
        sections = self._identify_structure(full_text, pages)
        
        return ParsedDocument(
            metadata=metadata,
            full_text=full_text.strip(),
            sections=sections,
            pages=pages
        )
    
    def _identify_structure(self, text: str, pages: List[str]) -> List[DocumentSection]:
        """Identifica estrutura do documento (seções, artigos, parágrafos)"""
        sections: List[DocumentSection] = []
        lines = text.split('\n')
        
        current_section: Optional[DocumentSection] = None
        current_content: List[str] = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Identificar títulos/seções (padrões comuns)
            # Ex: "SEÇÃO 5.2.3", "Capítulo 2", "5.2.3 Validação"
            section_match = re.match(r'^(?:SEÇÃO|Seção|CAPÍTULO|Capítulo|CAPITULO)\s+([\d.]+)', line, re.IGNORECASE)
            if section_match:
                # Salvar seção anterior
                if current_section:
                    current_section.content = '\n'.join(current_content)
                    sections.append(current_section)
                
                # Nova seção
                current_section = DocumentSection(
                    title=line,
                    level=1,
                    content=""
                )
                current_content = []
                continue
            
            # Identificar artigos (ex: "Art. 12", "Artigo 12")
            article_match = re.match(r'^(?:Art\.?|Artigo)\s+(\d+)', line, re.IGNORECASE)
            if article_match:
                if current_section:
                    current_section.content = '\n'.join(current_content)
                    sections.append(current_section)
                
                current_section = DocumentSection(
                    title=line,
                    level=2,
                    content="",
                    article_number=article_match.group(1)
                )
                current_content = []
                continue
            
            # Identificar parágrafos (ex: "§ 3º", "Parágrafo único")
            paragraph_match = re.match(r'^§\s*(\d+|único|ÚNICO)', line, re.IGNORECASE)
            if paragraph_match:
                if current_section:
                    current_section.content = '\n'.join(current_content)
                    sections.append(current_section)
                
                current_section = DocumentSection(
                    title=line,
                    level=3,
                    content="",
                    article_number=paragraph_match.group(1)
                )
                current_content = []
                continue
            
            # Adicionar conteúdo à seção atual
            if current_section:
                current_content.append(line)
            else:
                # Criar seção genérica se não houver uma
                current_section = DocumentSection(
                    title="Introdução",
                    level=1,
                    content=""
                )
                current_content = [line]
        
        # Salvar última seção
        if current_section:
            current_section.content = '\n'.join(current_content)
            sections.append(current_section)
        
        return sections


class DOCXParser:
    """Parser para documentos DOCX"""
    
    def __init__(self, file_path: str):
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")
        
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx não está instalado")
    
    def calculate_hash(self) -> str:
        """Calcula hash SHA-256 do arquivo"""
        sha256_hash = hashlib.sha256()
        with open(self.file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    
    def extract_metadata(self) -> DocumentMetadata:
        """Extrai metadados do DOCX e do nome do arquivo"""
        # Extrair metadados do nome do arquivo (mesma lógica do PDF)
        filename = self.file_path.stem
        metadata = PDFParser._parse_filename(filename)  # Reutilizar método estático
        
        # Calcular hash
        metadata.hash_arquivo = self.calculate_hash()
        metadata.arquivo_path = str(self.file_path)
        
        # Tentar extrair metadados do DOCX
        try:
            doc = DocxDocument(self.file_path)
            core_props = doc.core_properties
            
            if not metadata.autor and core_props.author:
                metadata.autor = core_props.author
            if not metadata.data_publicacao and core_props.created:
                metadata.data_publicacao = core_props.created.date()
            
            metadata.metadata['docx_metadata'] = {
                'title': core_props.title,
                'subject': core_props.subject,
                'creator': core_props.creator,
                'last_modified_by': core_props.last_modified_by,
                'created': str(core_props.created) if core_props.created else None,
                'modified': str(core_props.modified) if core_props.modified else None,
            }
        except Exception as e:
            print(f"⚠️  Erro ao extrair metadados do DOCX: {e}")
        
        return metadata
    
    def extract_text(self) -> ParsedDocument:
        """Extrai texto completo e estrutura do DOCX"""
        metadata = self.extract_metadata()
        full_text = ""
        sections: List[DocumentSection] = []
        
        doc = DocxDocument(self.file_path)
        
        # Extrair texto completo
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
                full_text += para.text + "\n"
        
        # Identificar estrutura baseada em estilos e formatação
        sections = self._identify_structure(doc)
        
        return ParsedDocument(
            metadata=metadata,
            full_text=full_text.strip(),
            sections=sections,
            pages=[]  # DOCX não tem páginas definidas
        )
    
    def _identify_structure(self, doc: DocxDocument) -> List[DocumentSection]:
        """Identifica estrutura do documento baseada em estilos"""
        sections: List[DocumentSection] = []
        current_section: Optional[DocumentSection] = None
        current_content: List[str] = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Verificar estilo do parágrafo
            style_name = para.style.name if para.style else ""
            
            # Títulos geralmente têm estilos como "Heading 1", "Heading 2", etc.
            if "Heading" in style_name or "Título" in style_name:
                level = 1
                if "1" in style_name:
                    level = 1
                elif "2" in style_name:
                    level = 2
                elif "3" in style_name:
                    level = 3
                
                # Salvar seção anterior
                if current_section:
                    current_section.content = '\n'.join(current_content)
                    sections.append(current_section)
                
                # Nova seção
                current_section = DocumentSection(
                    title=text,
                    level=level,
                    content=""
                )
                current_content = []
            else:
                # Conteúdo normal
                if current_section:
                    current_content.append(text)
                else:
                    # Criar seção genérica
                    current_section = DocumentSection(
                        title="Introdução",
                        level=1,
                        content=""
                    )
                    current_content = [text]
        
        # Salvar última seção
        if current_section:
            current_section.content = '\n'.join(current_content)
            sections.append(current_section)
        
        return sections


def parse_document(file_path: str) -> ParsedDocument:
    """
    Função principal para parsear um documento (PDF ou DOCX)
    
    Args:
        file_path: Caminho para o arquivo PDF ou DOCX
        
    Returns:
        ParsedDocument com texto, estrutura e metadados
    """
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")
    
    suffix = path.suffix.lower()
    
    if suffix == '.pdf':
        parser = PDFParser(file_path)
        return parser.extract_text()
    elif suffix in ['.docx', '.doc']:
        if suffix == '.doc':
            raise ValueError("Arquivos .doc (não .docx) não são suportados. Converta para .docx primeiro.")
        parser = DOCXParser(file_path)
        return parser.extract_text()
    else:
        raise ValueError(f"Formato não suportado: {suffix}. Use .pdf ou .docx")


if __name__ == "__main__":
    # Teste básico
    import sys
    
    if len(sys.argv) < 2:
        print("Uso: python document_parser.py <caminho_do_arquivo>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    try:
        print(f"📄 Parseando documento: {file_path}")
        parsed = parse_document(file_path)
        
        print(f"\n✅ Documento parseado com sucesso!")
        print(f"   Tipo: {parsed.metadata.documento_tipo}")
        print(f"   Nome: {parsed.metadata.documento_nome}")
        print(f"   Versão: {parsed.metadata.versao}")
        print(f"   Hash: {parsed.metadata.hash_arquivo[:16]}...")
        print(f"   Seções encontradas: {len(parsed.sections)}")
        print(f"   Tamanho do texto: {len(parsed.full_text)} caracteres")
        
        if parsed.sections:
            print(f"\n📑 Primeiras seções:")
            for i, section in enumerate(parsed.sections[:5], 1):
                print(f"   {i}. {section.title[:60]}...")
        
    except Exception as e:
        print(f"❌ Erro: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

